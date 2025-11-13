# app.py
"""
AI Research Paper Assistant (Full-featured)
- Upload PDFs, extract text
- Multi-level summarization via Hugging Face Inference API
- Citation & reference extraction (regex heuristics)
- Semantic search using sentence-transformers (all-MiniLM-L6-v2)
- Topic clustering with KMeans
- Paper-to-paper recommendations using embedding similarity
- Export summaries as .docx, .pdf, .csv
"""

import os
import time
import json
import re
from io import BytesIO
from dataclasses import dataclass, asdict
from typing import List, Dict, Any, Optional

import streamlit as st
import PyPDF2
import requests
from sentence_transformers import SentenceTransformer, util
from sklearn.cluster import KMeans
import numpy as np
import pandas as pd
from docx import Document
from fpdf2 import FPDF
import matplotlib.pyplot as plt
from dotenv import load_dotenv

# -------------------------
#  User-provided HF key (from your message)
#  You can remove this and rely on .env or environment variables instead.
USER_PROVIDED_HF_KEY = "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"
# -------------------------

# Load .env if present
load_dotenv()
# Prefer environment variable; if not present, use the provided key as fallback
if not os.getenv("HUGGINGFACE_API_KEY") and USER_PROVIDED_HF_KEY:
    os.environ["HUGGINGFACE_API_KEY"] = USER_PROVIDED_HF_KEY

# ------------------ Configuration / Defaults ------------------
HF_API_URL_BASE = "https://api-inference.huggingface.co/models"
HF_SUMMARY_MODEL_DEFAULT = "facebook/bart-large-cnn"
EMBEDDING_MODEL_DEFAULT = "sentence-transformers/all-MiniLM-L6-v2"

# ------------------ Data Model ------------------
@dataclass
class ResearchPaper:
    id: str
    filename: str
    status: str  # "processing" or "completed"
    title: str = ""
    authors: List[str] = None
    abstract: str = ""
    executive_summary: str = ""
    key_findings: List[str] = None
    methodology: str = ""
    sections: List[Dict[str, str]] = None
    raw_text: str = ""
    embedding: Optional[np.ndarray] = None
    topics: List[str] = None

    def to_dict(self):
        d = asdict(self)
        if self.embedding is not None:
            d["embedding"] = self.embedding.tolist()
        return d

# ------------------ Utilities ------------------
def make_id(filename: str) -> str:
    return f"{int(time.time()*1000)}_{filename}"

def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    reader = PyPDF2.PdfReader(BytesIO(pdf_bytes))
    texts = []
    for page in reader.pages:
        try:
            txt = page.extract_text()
            if txt:
                texts.append(txt)
        except Exception:
            continue
    return "\n\n".join(texts).strip()

# ------------------ Hugging Face Inference API ------------------
def hf_call_inference(model: str, inputs: str, params: Dict[str, Any] = None, timeout: int = 120) -> str:
    api_key = os.getenv("HUGGINGFACE_API_KEY")
    if not api_key:
        raise RuntimeError("Missing HUGGINGFACE_API_KEY environment variable.")
    headers = {"Authorization": f"Bearer {api_key}"}
    url = f"{HF_API_URL_BASE}/{model}"
    payload = {"inputs": inputs}
    if params:
        payload["parameters"] = params
    resp = requests.post(url, headers=headers, json=payload, timeout=timeout)
    if resp.status_code != 200:
        raise RuntimeError(f"Hugging Face inference API error {resp.status_code}: {resp.text}")
    data = resp.json()
    # Common formats: list with 'generated_text' or dict with 'summary_text'
    if isinstance(data, list) and data and "generated_text" in data[0]:
        return data[0]["generated_text"]
    if isinstance(data, dict) and "summary_text" in data:
        return data["summary_text"]
    # fallback to str
        # Check if API returned an error message
    if isinstance(data, dict) and "error" in data:
        raise RuntimeError(f"HF API error: {data['error']}")
    # Check if model is loading
    if isinstance(data, dict) and "estimated_time" in data:
        raise RuntimeError("Model is loading, please try again in a moment")
    return json.dumps(data)

def hf_summarize(text: str, model: str, length: str = "medium") -> str:
    if not text.strip():
        return ""
    # mapping lengths to model params (for summarization models that accept min/max length)
    if length == "short":
        params = {"min_length": 20, "max_length": 80}
    elif length == "long":
        params = {"min_length": 150, "max_length": 600}
    else:
        params = {"min_length": 80, "max_length": 200}
    prompt = text if len(text) < 20000 else text[:20000]
    return hf_call_inference(model, prompt, params=params)

# ------------------ Citation & Reference Extraction ------------------
CITATION_PATTERNS = [
    r"\[\d+\]",                         # [1], [23]
    r"\([A-Z][A-Za-z\-]+ et al\., \d{4}\)",  # (Smith et al., 2020)
    r"\([A-Z][A-Za-z\-]+, \d{4}\)",     # (Smith, 2020)
    r"\b[A-Z][A-Za-z\-]+ et al\.\b",    # Smith et al.
    r"\bdoi:\s*10\.\d{4,9}/[-._;()/:A-Z0-9]+",  # doi:...
]
def extract_citations(text: str) -> List[str]:
    found = set()
    for pat in CITATION_PATTERNS:
        for m in re.findall(pat, text):
            found.add(m)
    refs = []
    lower = text.lower()
    if "references" in lower:
        try:
            idx = lower.index("references")
            refs_text = text[idx:]
            for line in refs_text.splitlines():
                line = line.strip()
                if not line: 
                    continue
                if re.match(r"^\[?\d+\]?", line) or "doi" in line.lower() or re.search(r"\d{4}", line):
                    refs.append(line)
            if refs:
                return list(dict.fromkeys(refs))[:50]
        except Exception:
            pass
    return sorted(found)

# ------------------ Embeddings, Search & Clustering ------------------
@st.cache_resource
def load_embedding_model(name: str):
    return SentenceTransformer(name)

def compute_embeddings_for_papers(model, papers: List[ResearchPaper]):
    texts = []
    ids = []
    for p in papers:
        txt = p.executive_summary or p.abstract or p.raw_text
        texts.append(txt if txt else p.filename)
        ids.append(p.id)
    if not texts:
        return
    embeddings = model.encode(texts, convert_to_numpy=True, show_progress_bar=False)
    for i, p in enumerate(papers):
        p.embedding = embeddings[i]

def semantic_search(model, papers: List[ResearchPaper], query: str, top_k: int = 5):
    q_emb = model.encode(query, convert_to_numpy=True)
    candidates = []
    for p in papers:
        if p.embedding is None:
            continue
        score = util.cos_sim(q_emb, p.embedding).item()
        candidates.append((score, p))
    candidates.sort(key=lambda x: x[0], reverse=True)
    return candidates[:top_k]

def cluster_papers(papers: List[ResearchPaper], n_clusters: int = 4):
    X = []
    ids = []
    for p in papers:
        if p.embedding is not None:
            X.append(p.embedding)
            ids.append(p.id)
    if not X:
        return {}
    X = np.vstack(X)
    kmeans = KMeans(n_clusters=min(n_clusters, len(X)), random_state=42).fit(X)
    clusters = {}
    for idx, label in enumerate(kmeans.labels_):
        pid = ids[idx]
        clusters.setdefault(int(label), []).append(pid)
    topics = {}
    for label, pids in clusters.items():
        words = []
        for pid in pids:
            p = next(p for p in papers if p.id == pid)
            txt = (p.executive_summary or p.abstract or "").lower()
            words += re.findall(r"\b[a-z]{4,}\b", txt)
        top = [w for w, c in pd.Series(words).value_counts().head(5).items()]
        topics[label] = top
    return {"clusters": clusters, "topics": topics}

# ------------------ Exports ------------------
def make_docx_summary(paper: ResearchPaper) -> bytes:
    doc = Document()
    doc.add_heading(paper.title or paper.filename, level=1)
    if paper.authors:
        doc.add_paragraph("Authors: " + ", ".join(paper.authors))
    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(paper.executive_summary or paper.abstract or "")
    doc.add_heading("Key Findings", level=2)
    if paper.key_findings:
        for k in paper.key_findings:
            doc.add_paragraph(k, style="List Bullet")
    doc.add_heading("Methodology", level=2)
    doc.add_paragraph(paper.methodology or "")
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def make_pdf_summary(paper: ResearchPaper) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", 'B', 14)
    pdf.multi_cell(0, 8, paper.title or paper.filename)
    pdf.ln(4)
    pdf.set_font("Arial", size=11)
    pdf.multi_cell(0, 7, "Executive Summary:\n" + (paper.executive_summary or paper.abstract or ""))
    pdf.ln(4)
    if paper.key_findings:
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 6, "Key Findings:")
        pdf.ln(6)
        pdf.set_font("Arial", size=11)
        for k in paper.key_findings:
            pdf.multi_cell(0, 6, "- " + k)
    return pdf.output()

def make_csv_export(papers: List[ResearchPaper]) -> bytes:
    rows = []
    for p in papers:
        rows.append({
            "id": p.id,
            "filename": p.filename,
            "title": p.title,
            "authors": "; ".join(p.authors or []),
            "abstract": p.abstract,
            "executive_summary": p.executive_summary,
            "key_findings": "|".join(p.key_findings or []),
            "methodology": p.methodology
        })
    df = pd.DataFrame(rows)
    return df.to_csv(index=False).encode('utf-8')

# ------------------ Processing Pipeline ------------------
def process_paper_bytes(file_bytes: bytes, filename: str, hf_summary_model: str, hf_summary_level: str = "medium") -> ResearchPaper:
    rp = ResearchPaper(id=make_id(filename), filename=filename, status="processing", authors=[], key_findings=[], sections=[], topics=[])
    text = extract_text_from_pdf_bytes(file_bytes)
    rp.raw_text = text
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    rp.title = lines[0] if lines else filename
    lower = text.lower()
    abstr = ""
    if "abstract" in lower:
        try:
            idx = lower.index("abstract")
            snippet = text[idx: idx + 3000]
            abstr = re.split(r"\n\s*\n", snippet, maxsplit=1)[0]
        except Exception:
            abstr = ""
    rp.abstract = abstr or (lines[1] if len(lines) > 1 else "")
    # Summarize with HF
    try:
        exec_sum = hf_summarize(text, hf_summary_model, length=hf_summary_level)
        rp.executive_summary = exec_sum
        sents = re.split(r'\.\s+', exec_sum.strip())
        rp.key_findings = [s.strip() for s in sents if s][:6]
    except Exception as e:
        rp.executive_summary = (rp.abstract or "")[:1000]
        rp.key_findings = []
    # Methodology extraction heuristics
    method = ""
            st.error(f"Summary generation failed: {str(e)}")
    for marker in ["methodology", "methods", "experimental setup", "materials and methods"]:
        lm = lower.find(marker)
        if lm != -1:
            snippet = text[lm:lm+2000]
            method = re.split(r"\n\s*\n", snippet, maxsplit=1)[0]
            break
    rp.methodology = method
    rp.status = "completed"
    return rp

# ------------------ Streamlit UI ------------------
st.set_page_config(page_title="AI Research Paper Assistant", layout="wide")
st.title("AI Research Paper Assistant (Hugging Face + Embeddings)")

st.sidebar.header("Navigation")
page = st.sidebar.selectbox("Choose page", ["Upload", "Library", "Search", "Clustering", "Compare & Export", "Settings", "Stats"])

if "papers" not in st.session_state:
    st.session_state.papers: Dict[str, ResearchPaper] = {}
if "embedding_model" not in st.session_state:
    st.session_state.embedding_model = None
if "last_cluster_result" not in st.session_state:
    st.session_state.last_cluster_result = None

# Settings panel values (persist across interactions)
if "hf_summary_model" not in st.session_state:
    st.session_state.hf_summary_model = HF_SUMMARY_MODEL_DEFAULT
if "embedding_model_name" not in st.session_state:
    st.session_state.embedding_model_name = EMBEDDING_MODEL_DEFAULT

with st.sidebar.expander("Models / Settings", expanded=False):
    st.session_state.hf_summary_model = st.text_input("HF summary model", value=st.session_state.hf_summary_model)
    st.session_state.embedding_model_name = st.text_input("Embedding model", value=st.session_state.embedding_model_name)
    st.write("Make sure HUGGINGFACE_API_KEY is set (or the provided key is loaded).")

# Upload Page
if page == "Upload":
    st.header("Upload PDFs")
    uploaded = st.file_uploader("Upload one or more PDF files", type=["pdf"], accept_multiple_files=True)
    sum_level = st.radio("Summary level for processing", ["short", "medium", "long"], index=1, horizontal=True)
    if uploaded and st.button("Process uploaded PDFs"):
        progress = st.progress(0)
        total = len(uploaded)
        for i, f in enumerate(uploaded):
            progress.progress(int((i / total) * 100))
            bytes_data = f.read()
            rp = process_paper_bytes(bytes_data, f.name, hf_summary_model=st.session_state.hf_summary_model, hf_summary_level=sum_level)
            st.session_state.papers[rp.id] = rp
            progress.progress(int(((i+1) / total) * 100))
        st.success("Processing complete.")
        # Load embedding model
        st.session_state.embedding_model = load_embedding_model(st.session_state.embedding_model_name)
        compute_embeddings_for_papers(st.session_state.embedding_model, list(st.session_state.papers.values()))

# Library Page
elif page == "Library":
    st.header("Library")
    papers = list(st.session_state.papers.values())
    if not papers:
        st.info("No papers in library yet. Upload on the Upload page.")
    else:
        for p in papers:
            with st.expander(f"{p.title or p.filename} — {p.status}", expanded=False):
                st.write("Filename:", p.filename)
                st.write("Executive Summary:")
                st.write(p.executive_summary or p.abstract or "—")
                st.write("Key findings:")
                for k in (p.key_findings or []):
                    st.write("- " + k)
                citations = extract_citations(p.raw_text)
                if citations:
                    st.write("Citations / References (detected):")
                    for c in citations[:20]:
                        st.write("- " + c)
                col1, col2, col3 = st.columns([1,1,1])
                with col1:
                    if st.button("Download .docx", key=f"docx_{p.id}"):
                        data = make_docx_summary(p)
                        st.download_button("Download DOCX", data, file_name=f"{p.filename}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                with col2:
                    if st.button("Download .pdf", key=f"pdf_{p.id}"):
                        data = make_pdf_summary(p)
                        st.download_button("Download PDF", data, file_name=f"{p.filename}.pdf", mime="application/pdf")
                with col3:
                    if st.button("View Raw", key=f"raw_{p.id}"):
                        st.text_area("Raw text", p.raw_text[:60000], height=300)

# Search Page
elif page == "Search":
    st.header("Semantic Search")
    if not st.session_state.papers:
        st.info("No papers uploaded.")
    else:
        if st.session_state.embedding_model is None:
            st.session_state.embedding_model = load_embedding_model(st.session_state.embedding_model_name)
            compute_embeddings_for_papers(st.session_state.embedding_model, list(st.session_state.papers.values()))
        query = st.text_input("Enter search query (topic, phrase)")
        k = st.slider("Top K results", 1, 10, 5)
        if st.button("Search") and query.strip():
            results = semantic_search(st.session_state.embedding_model, list(st.session_state.papers.values()), query, top_k=k)
            if not results:
                st.write("No results (embeddings may be missing).")
            else:
                for score, p in results:
                    st.markdown(f"### {p.title} — (score: {score:.3f})")
                    st.write(p.executive_summary[:500] + ("..." if len(p.executive_summary) > 500 else ""))

# Clustering Page
elif page == "Clustering":
    st.header("Topic Clustering")
    papers = list(st.session_state.papers.values())
    if not papers:
        st.info("No papers uploaded.")
    else:
        if st.session_state.embedding_model is None:
            st.session_state.embedding_model = load_embedding_model(st.session_state.embedding_model_name)
        compute_embeddings_for_papers(st.session_state.embedding_model, papers)
        n_clusters = st.slider("Number of clusters", 2, min(10, max(2, len(papers))), 4)
        if st.button("Run Clustering"):
            result = cluster_papers(papers, n_clusters=n_clusters)
            st.session_state.last_cluster_result = result
            clusters = result.get("clusters", {})
            topics = result.get("topics", {})
            st.success(f"Clustering produced {len(clusters)} clusters.")
            for label, pids in clusters.items():
                st.markdown(f"#### Cluster {label} — topic keywords: {', '.join(topics.get(label, []))}")
                for pid in pids:
                    p = next(p for p in papers if p.id == pid)
                    st.write("- " + (p.title or p.filename))
            sizes = {label: len(v) for label, v in clusters.items()}
            df = pd.DataFrame({"cluster": list(sizes.keys()), "size": list(sizes.values())})
            st.bar_chart(df.set_index("cluster"))

# Compare & Export Page
elif page == "Compare & Export":
    st.header("Compare Papers & Export")
    papers = list(st.session_state.papers.values())
    if len(papers) < 1:
        st.info("Upload papers first.")
    else:
        choices = {p.id: p.filename for p in papers}
        selected = st.multiselect("Select papers to include", options=list(choices.keys()), format_func=lambda x: choices[x])
        if selected:
            sel_papers = [st.session_state.papers[sid] for sid in selected]
            if st.button("Show recommendations (by similarity)"):
                if st.session_state.embedding_model is None:
                    st.session_state.embedding_model = load_embedding_model(st.session_state.embedding_model_name)
                    compute_embeddings_for_papers(st.session_state.embedding_model, papers)
                for p in sel_papers:
                    st.markdown(f"### Recommendations for: {p.title or p.filename}")
                    sims = []
                    for other in papers:
                        if other.id == p.id or other.embedding is None or p.embedding is None:
                            continue
                        score = util.cos_sim(p.embedding, other.embedding).item()
                        sims.append((score, other))
                    sims = sorted(sims, key=lambda x: x[0], reverse=True)[:5]
                    for score, other in sims:
                        st.write(f"- {other.title or other.filename} (score: {score:.3f})")
            if st.button("Export selected as CSV"):
                csv_bytes = make_csv_export(sel_papers)
                st.download_button("Download CSV", csv_bytes, file_name="papers_export.csv", mime="text/csv")
            if st.button("Export selected as DOCX bundle"):
                doc = Document()
                for p in sel_papers:
                    doc.add_heading(p.title or p.filename, level=1)
                    if p.authors:
                        doc.add_paragraph("Authors: " + ", ".join(p.authors))
                    doc.add_heading("Executive Summary", level=2)
                    doc.add_paragraph(p.executive_summary or p.abstract or "")
                    doc.add_heading("Key Findings", level=2)
                    for k in (p.key_findings or []):
                        doc.add_paragraph(k, style="List Bullet")
                    doc.add_page_break()
                bio = BytesIO()
                doc.save(bio)
                bio.seek(0)
                st.download_button("Download combined DOCX", bio.read(), file_name="combined_summaries.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# Settings Page
elif page == "Settings":
    st.header("Settings & Diagnostics")
    st.write("Hugging Face inference model:", st.session_state.hf_summary_model)
    st.write("Embedding model:", st.session_state.embedding_model_name)
    hf_key = os.getenv("HUGGINGFACE_API_KEY")
    if hf_key:
        st.success("HUGGINGFACE_API_KEY detected in environment.")
    else:
        st.warning("HUGGINGFACE_API_KEY not found. Please ensure the key is set.")
    if st.button("Reload embedding model"):
        st.session_state.embedding_model = load_embedding_model(st.session_state.embedding_model_name)
        st.success("Embedding model loaded.")

# Stats Page
elif page == "Stats":
    st.header("Stats & Quick Insights")
    papers = list(st.session_state.papers.values())
    st.metric("Total papers", len(papers))
    st.metric("Papers with embeddings", sum(1 for p in papers if p.embedding is not None))
    st.metric("Total key findings extracted", sum(len(p.key_findings or []) for p in papers))
    words = []
    for p in papers:
        s = (p.executive_summary or p.abstract or "").lower()
        words += re.findall(r"\b[a-z]{4,}\b", s)
    if words:
        counts = pd.Series(words).value_counts().head(40)
        fig, ax = plt.subplots(figsize=(8,4))
        counts.plot.bar(ax=ax)
        st.pyplot(fig)

st.sidebar.markdown("---")
st.sidebar.write("Built with Streamlit • Hugging Face • sentence-transformers")
